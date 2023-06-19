from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
import os

def searchFunc(text, query):
    newList = []
    start = 0
    text = text.lower()

    while(start != -1):
        start = text.find(query, start)
        if(start == -1):
            break
        newList.append(start)
        start = start + 1
    return newList

def combineArrays(bigArr, insertArr, insertArrNum):
    if len(bigArr) == 0:
        for i in range(len(insertArr)):
            bigArr.append([insertArr[i], insertArrNum])
        return bigArr
    
    size_1 = len(bigArr)
    size_2 = len(insertArr)
    
    res = []
    i, j = 0, 0
    
    while i < size_1 and j < size_2:
        if bigArr[i][0] < insertArr[j]:
            res.append(bigArr[i])
            i += 1
    
        else:
            res.append([insertArr[j], insertArrNum])
            j += 1
    for thing in insertArr[j:]:
        res.append([thing, insertArrNum])
    res += bigArr[i:]
    return res

#Grrrr
def colorFind(query):
    match query:
        case "blue":
            return WD_COLOR_INDEX.BLUE
        case "bright green":
            return WD_COLOR_INDEX.BRIGHT_GREEN
        case "dark blue":
            return WD_COLOR_INDEX.DARK_BLUE
        case "dark red":
            return WD_COLOR_INDEX.DARK_RED
        case "dark yellow":
            return WD_COLOR_INDEX.DARK_YELLOW
        case "green":
            return WD_COLOR_INDEX.GREEN
        case "pink":
            return WD_COLOR_INDEX.PINK
        case "red":
            return WD_COLOR_INDEX.RED
        case "teal":
            return WD_COLOR_INDEX.TEAL
        case "turquoise":
            return WD_COLOR_INDEX.TURQUOISE
        case "violet":
            return WD_COLOR_INDEX.VIOLET
        case _:
            print(query)
            return WD_COLOR_INDEX.YELLOW

def highlightWord(path, queryList, colorList):
    document = Document(path)

    for paragraph in document.paragraphs:
        bigList = []
        paragraphText = paragraph.text
        for i in range(len(queryList)):
            returnList = searchFunc(paragraphText, queryList[i])
            if len(returnList) != 0:
                bigList = combineArrays(bigList, returnList, i)
        
        if(len(bigList) > 0):
            paragraph.clear()
            start = 0
            for word in bigList:
                paragraph.add_run(paragraphText[start:word[0]])
                paragraph.add_run(queryList[word[1]]).font.highlight_color = colorList[word[1]]
                start = word[0] + len(queryList[word[1]])
            paragraph.add_run(paragraphText[start:])
            paragraph.style.font.name = "Arial"
            paragraph.style.font.size = Pt(10)

    copyName = path[:len(path) - 5] + " - Copy.DOCX"
    document.save(copyName)

def runOnAllFilesInDirectory(path, queryList, colorList):
    dirList = os.listdir(path)
    #print(dirList)
    for dir in dirList:
        if os.path.isdir(os.path.join(path, dir)):
            runOnAllFilesInDirectory(os.path.join(path, dir), queryList, colorList)
        else:
            if dir.endswith(".DOCX"):
                highlightWord(os.path.join(path, dir), queryList, colorList)
                print(dir)
    return 0

def deleteCopies(path):
    dirList = os.listdir(path)
    #print(dirList)
    for dir in dirList:
        if os.path.isdir(os.path.join(path, dir)):
            deleteCopies(os.path.join(path, dir))
        else:
            if dir.endswith(" - Copy.DOCX"):
                os.remove(os.path.join(path, dir))
                print(dir)
    return 0

def main():
    print("Hello World")
    queryFile = open("query.txt", "r")
    queryList = queryFile.readlines()
    queryFile.close()

    colorList = []
    for i in range(len(queryList)):
        temp = queryList[i].split(", ")
        queryList[i] = temp[0].lower()
        if(len(temp) < 2):
            colorList.append(WD_COLOR_INDEX.YELLOW)
        else:
            colorList.append(colorFind(temp[1].strip().lower()))
    print(colorList)

    print(queryList)

    file = os.path.join('Raw Data')
    deleteCopies(file)
    runOnAllFilesInDirectory(file, queryList, colorList)
    

if(__name__ == "__main__"):
    main()